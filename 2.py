#beta2
# -*- coding: utf-8 -*-
"""
AB 词汇/句子对照记忆（GUI）
- 读取：Excel(xlsx/xls)、CSV/TSV/TXT、Word(docx表格)
- 自动建立 A-B 对应
- 练习：给x选y（选择题）、给x填y（填空）
- 方向：A->B 或 B->A（x、y可互换）
- GUI：全程鼠标可点（填空需要键盘输入答案）
"""

import os
import random
import tkinter as tk
from tkinter import ttk, filedialog, messagebox

# --- optional deps (按需安装) ---
# pandas + openpyxl: 读Excel/CSV
# python-docx: 读docx表格

def _try_import_pandas():
    try:
        import pandas as pd
        return pd
    except Exception:
        return None

def _try_import_docx():
    try:
        from docx import Document
        return Document
    except Exception:
        return None

def normalize_text(s: str, case_insensitive: bool = True) -> str:
    if s is None:
        return ""
    s = str(s).strip()
    # 常见的“全角空格”
    s = s.replace("\u3000", " ").strip()
    if case_insensitive:
        s = s.lower()
    return s

class Deck:
    def __init__(self):
        self.pairs = []   # list[(A,B)]
        self.source = ""  # file path
        self.headers = [] # list[str]
        self.preview_rows = []  # list[list[str]]

    def set_pairs(self, pairs, headers=None, preview_rows=None, source=""):
        # 清洗：去掉空行、重复行
        cleaned = []
        seen = set()
        for a, b in pairs:
            a2 = str(a).strip() if a is not None else ""
            b2 = str(b).strip() if b is not None else ""
            if not a2 and not b2:
                continue
            key = (a2, b2)
            if key in seen:
                continue
            seen.add(key)
            cleaned.append((a2, b2))

        self.pairs = cleaned
        self.headers = headers or []
        self.preview_rows = preview_rows or []
        self.source = source

    def size(self):
        return len(self.pairs)

class ABApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("A-B 对照记忆（GUI）")
        self.geometry("980x620")
        self.minsize(900, 560)

        self.deck = Deck()

        # quiz state
        self.quiz_mode = tk.StringVar(value="mc")  # mc / fill
        self.direction = tk.StringVar(value="A2B") # A2B / B2A
        self.shuffle = tk.BooleanVar(value=True)
        self.case_insensitive = tk.BooleanVar(value=True)
        self.num_choices = tk.IntVar(value=4)

        self.current_index = -1
        self.order = []
        self.correct = 0
        self.total = 0
        self.current_x = ""
        self.current_y = ""
        self.current_options = []

        # load state
        self.file_path = tk.StringVar(value="")
        self.sheet_name = tk.StringVar(value="")
        self.has_header = tk.BooleanVar(value=True)
        self.col_a = tk.IntVar(value=0)
        self.col_b = tk.IntVar(value=1)

        self._build_ui()

    # ---------------- UI ----------------
    def _build_ui(self):
        self.style = ttk.Style()
        # macOS 下有时默认主题较怪，尽量保持系统风格
        try:
            self.style.theme_use("clam")
        except Exception:
            pass

        nb = ttk.Notebook(self)
        nb.pack(fill="both", expand=True)

        self.tab_load = ttk.Frame(nb)
        self.tab_setup = ttk.Frame(nb)
        self.tab_quiz = ttk.Frame(nb)

        nb.add(self.tab_load, text="1) 导入表格")
        nb.add(self.tab_setup, text="2) 设置练习")
        nb.add(self.tab_quiz, text="3) 开始练习")

        self._build_tab_load()
        self._build_tab_setup()
        self._build_tab_quiz()

    def _build_tab_load(self):
        frm = self.tab_load
        frm.columnconfigure(0, weight=1)
        frm.rowconfigure(3, weight=1)

        top = ttk.Frame(frm)
        top.grid(row=0, column=0, sticky="ew", padx=12, pady=12)
        top.columnconfigure(1, weight=1)

        ttk.Button(top, text="选择文件…", command=self.pick_file).grid(row=0, column=0, padx=(0,8))
        ttk.Entry(top, textvariable=self.file_path).grid(row=0, column=1, sticky="ew")

        opt = ttk.LabelFrame(frm, text="导入选项")
        opt.grid(row=1, column=0, sticky="ew", padx=12, pady=(0,12))
        opt.columnconfigure(5, weight=1)

        # header checkbox
        ttk.Checkbutton(opt, text="第一行是表头", variable=self.has_header, command=self.refresh_preview).grid(row=0, column=0, padx=8, pady=8, sticky="w")

        # sheet (for excel)
        ttk.Label(opt, text="工作表:").grid(row=0, column=1, padx=(8,2), pady=8, sticky="e")
        self.sheet_combo = ttk.Combobox(opt, textvariable=self.sheet_name, state="disabled", width=22)
        self.sheet_combo.grid(row=0, column=2, padx=(2,8), pady=8, sticky="w")
        self.sheet_combo.bind("<<ComboboxSelected>>", lambda e: self.refresh_preview())

        # columns
        ttk.Label(opt, text="列 A(题干):").grid(row=0, column=3, padx=(8,2), pady=8, sticky="e")
        self.col_a_spin = ttk.Spinbox(opt, from_=0, to=999, textvariable=self.col_a, width=6, command=self.refresh_preview)
        self.col_a_spin.grid(row=0, column=4, padx=(2,8), pady=8, sticky="w")

        ttk.Label(opt, text="列 B(答案):").grid(row=0, column=5, padx=(8,2), pady=8, sticky="e")
        self.col_b_spin = ttk.Spinbox(opt, from_=0, to=999, textvariable=self.col_b, width=6, command=self.refresh_preview)
        self.col_b_spin.grid(row=0, column=6, padx=(2,8), pady=8, sticky="w")

        ttk.Button(opt, text="刷新预览", command=self.refresh_preview).grid(row=0, column=7, padx=8, pady=8)

        # info line
        self.load_info = tk.StringVar(value="未导入。请选择一个表格文件。")
        ttk.Label(frm, textvariable=self.load_info).grid(row=2, column=0, sticky="ew", padx=12, pady=(0,6))

        # preview
        preview_box = ttk.LabelFrame(frm, text="预览（前 30 行）")
        preview_box.grid(row=3, column=0, sticky="nsew", padx=12, pady=(0,12))
        preview_box.columnconfigure(0, weight=1)
        preview_box.rowconfigure(0, weight=1)

        self.preview = ttk.Treeview(preview_box, columns=("A","B"), show="headings")
        self.preview.heading("A", text="A（列n）")
        self.preview.heading("B", text="B（列n+1）")
        self.preview.column("A", width=420, anchor="w")
        self.preview.column("B", width=420, anchor="w")
        self.preview.grid(row=0, column=0, sticky="nsew")

        yscroll = ttk.Scrollbar(preview_box, orient="vertical", command=self.preview.yview)
        self.preview.configure(yscrollcommand=yscroll.set)
        yscroll.grid(row=0, column=1, sticky="ns")

        btns = ttk.Frame(frm)
        btns.grid(row=4, column=0, sticky="ew", padx=12, pady=(0,12))
        btns.columnconfigure(0, weight=1)

        ttk.Button(btns, text="建立 A-B 对应关系 ✅", command=self.build_pairs).pack(side="right")

    def _build_tab_setup(self):
        frm = self.tab_setup
        frm.columnconfigure(0, weight=1)

        box = ttk.LabelFrame(frm, text="练习设置")
        box.grid(row=0, column=0, sticky="ew", padx=12, pady=12)
        for i in range(6):
            box.columnconfigure(i, weight=1)

        # mode
        ttk.Label(box, text="模式:").grid(row=0, column=0, padx=8, pady=10, sticky="e")
        ttk.Radiobutton(box, text="给 x 选 y（选择题）", variable=self.quiz_mode, value="mc").grid(row=0, column=1, padx=8, pady=10, sticky="w")
        ttk.Radiobutton(box, text="给 x 填 y（填空）", variable=self.quiz_mode, value="fill").grid(row=0, column=2, padx=8, pady=10, sticky="w")

        # direction
        ttk.Label(box, text="方向:").grid(row=1, column=0, padx=8, pady=10, sticky="e")
        ttk.Radiobutton(box, text="A → B", variable=self.direction, value="A2B").grid(row=1, column=1, padx=8, pady=10, sticky="w")
        ttk.Radiobutton(box, text="B → A", variable=self.direction, value="B2A").grid(row=1, column=2, padx=8, pady=10, sticky="w")

        # options
        ttk.Checkbutton(box, text="随机打乱", variable=self.shuffle).grid(row=2, column=1, padx=8, pady=10, sticky="w")
        ttk.Checkbutton(box, text="忽略大小写（填空）", variable=self.case_insensitive).grid(row=2, column=2, padx=8, pady=10, sticky="w")

        ttk.Label(box, text="选择题选项数:").grid(row=2, column=0, padx=8, pady=10, sticky="e")
        self.choice_spin = ttk.Spinbox(box, from_=2, to=8, textvariable=self.num_choices, width=6)
        self.choice_spin.grid(row=2, column=3, padx=8, pady=10, sticky="w")

        hint = ttk.Label(frm, foreground="#444",
                         text="⚠️ 先去“导入表格”建立 A-B 对应关系；再来这里设置；最后去“开始练习”。")
        hint.grid(row=1, column=0, sticky="w", padx=12, pady=(0,12))

        btns = ttk.Frame(frm)
        btns.grid(row=2, column=0, sticky="ew", padx=12, pady=(0,12))
        btns.columnconfigure(0, weight=1)

        ttk.Button(btns, text="初始化题目顺序 ▶️", command=self.reset_quiz).pack(side="right")

    def _build_tab_quiz(self):
        frm = self.tab_quiz
        frm.columnconfigure(0, weight=1)
        frm.rowconfigure(2, weight=1)

        top = ttk.Frame(frm)
        top.grid(row=0, column=0, sticky="ew", padx=12, pady=12)
        top.columnconfigure(3, weight=1)

        self.score_var = tk.StringVar(value="得分：0 / 0")
        ttk.Label(top, textvariable=self.score_var).grid(row=0, column=0, padx=(0,12), sticky="w")

        self.progress_var = tk.StringVar(value="进度：- / -")
        ttk.Label(top, textvariable=self.progress_var).grid(row=0, column=1, padx=(0,12), sticky="w")

        ttk.Button(top, text="下一题 ▶️", command=self.next_question).grid(row=0, column=2, padx=(0,8))
        ttk.Button(top, text="重置 ♻️", command=self.reset_quiz).grid(row=0, column=3, sticky="e")

        qbox = ttk.LabelFrame(frm, text="题目（x）")
        qbox.grid(row=1, column=0, sticky="ew", padx=12, pady=(0,12))
        qbox.columnconfigure(0, weight=1)

        self.question_text = tk.StringVar(value="（尚未开始）")
        self.q_label = ttk.Label(qbox, textvariable=self.question_text, wraplength=920, font=("Arial", 16))
        self.q_label.grid(row=0, column=0, sticky="ew", padx=12, pady=12)

        abox = ttk.LabelFrame(frm, text="作答（y）")
        abox.grid(row=2, column=0, sticky="nsew", padx=12, pady=(0,12))
        abox.columnconfigure(0, weight=1)
        abox.rowconfigure(1, weight=1)

        self.answer_area = ttk.Frame(abox)
        self.answer_area.grid(row=0, column=0, sticky="ew", padx=12, pady=12)
        self.answer_area.columnconfigure(0, weight=1)

        # fill mode widgets
        self.fill_frame = ttk.Frame(self.answer_area)
        self.fill_frame.grid(row=0, column=0, sticky="ew")
        self.fill_frame.columnconfigure(1, weight=1)

        ttk.Label(self.fill_frame, text="请输入答案：").grid(row=0, column=0, padx=(0,8), pady=8, sticky="w")
        self.fill_entry = ttk.Entry(self.fill_frame)
        self.fill_entry.grid(row=0, column=1, padx=(0,8), pady=8, sticky="ew")
        ttk.Button(self.fill_frame, text="提交 ✅", command=self.submit_fill).grid(row=0, column=2, pady=8)

        # multiple choice widgets
        self.mc_frame = ttk.Frame(self.answer_area)
        self.mc_frame.grid(row=1, column=0, sticky="ew")
        self.mc_frame.columnconfigure(0, weight=1)

        self.mc_buttons = []
        for i in range(8):  # 最大 8 个选项
            btn = ttk.Button(self.mc_frame, text=f"选项{i+1}", command=lambda k=i: self.submit_choice(k))
            self.mc_buttons.append(btn)

        self.feedback_var = tk.StringVar(value="")
        self.feedback_label = ttk.Label(abox, textvariable=self.feedback_var, wraplength=920)
        self.feedback_label.grid(row=2, column=0, sticky="ew", padx=12, pady=(0,12))

        self.update_answer_widgets_visibility()

    # ---------------- Data Loading ----------------
    def pick_file(self):
        path = filedialog.askopenfilename(
            title="选择表格文件",
            filetypes=[
                ("All supported", "*.xlsx *.xls *.csv *.tsv *.txt *.docx"),
                ("Excel", "*.xlsx *.xls"),
                ("CSV/TSV/TXT", "*.csv *.tsv *.txt"),
                ("Word (tables)", "*.docx"),
                ("All files", "*.*"),
            ],
        )
        if not path:
            return
        self.file_path.set(path)
        self.deck = Deck()
        self.load_info.set("已选择文件，准备预览…")
        self._prepare_sheet_list()
        self.refresh_preview()

    def _prepare_sheet_list(self):
        path = self.file_path.get().strip()
        ext = os.path.splitext(path)[1].lower()

        if ext in (".xlsx", ".xls"):
            pd = _try_import_pandas()
            if pd is None:
                messagebox.showerror("缺少依赖", "读取 Excel 需要安装 pandas + openpyxl。\n\npip install pandas openpyxl")
                self.sheet_combo.configure(state="disabled", values=[])
                self.sheet_name.set("")
                return
            try:
                xls = pd.ExcelFile(path)
                sheets = xls.sheet_names
                self.sheet_combo.configure(state="readonly", values=sheets)
                if sheets:
                    self.sheet_name.set(sheets[0])
            except Exception as e:
                self.sheet_combo.configure(state="disabled", values=[])
                self.sheet_name.set("")
                self.load_info.set(f"无法读取 Excel：{e}")
        else:
            self.sheet_combo.configure(state="disabled", values=[])
            self.sheet_name.set("")

    def refresh_preview(self):
        # clear preview
        for item in self.preview.get_children():
            self.preview.delete(item)

        path = self.file_path.get().strip()
        if not path:
            self.load_info.set("未导入。请选择一个表格文件。")
            return

        try:
            rows, headers = self._load_table_preview(path, limit=30)
            # set headings with column indexes (user chooses by index)
            a_idx = int(self.col_a.get())
            b_idx = int(self.col_b.get())
            self.preview.heading("A", text=f"A（列 {a_idx}）")
            self.preview.heading("B", text=f"B（列 {b_idx}）")

            shown = 0
            for r in rows:
                a = r[a_idx] if a_idx < len(r) else ""
                b = r[b_idx] if b_idx < len(r) else ""
                self.preview.insert("", "end", values=(a, b))
                shown += 1

            self.load_info.set(f"预览成功：显示 {shown} 行。当前列选择 A={a_idx}, B={b_idx}。")
        except Exception as e:
            self.load_info.set(f"预览失败：{e}")

    def _load_table_preview(self, path, limit=30):
        ext = os.path.splitext(path)[1].lower()
        has_header = bool(self.has_header.get())

        if ext in (".xlsx", ".xls", ".csv", ".tsv", ".txt"):
            pd = _try_import_pandas()
            if pd is None:
                raise RuntimeError("缺少 pandas。请安装：pip install pandas openpyxl")

            if ext in (".xlsx", ".xls"):
                sheet = self.sheet_name.get().strip() or 0
                df = pd.read_excel(path, sheet_name=sheet, header=0 if has_header else None, dtype=str)
            else:
                # csv/tsv/txt: 尝试多编码
                sep = "\t" if ext == ".tsv" else ("," if ext == ".csv" else None)
                encodings = ["utf-8-sig", "utf-8", "gbk", "gb18030"]
                last_err = None
                for enc in encodings:
                    try:
                        df = pd.read_csv(path, sep=sep, header=0 if has_header else None, dtype=str, encoding=enc, engine="python")
                        last_err = None
                        break
                    except Exception as e:
                        last_err = e
                        df = None
                if df is None:
                    raise RuntimeError(f"读取文本表格失败：{last_err}")

            if has_header:
                headers = [str(c) for c in df.columns.tolist()]
                data = df.head(limit).fillna("").values.tolist()
                # 把每个单元转成字符串
                data = [[("" if v is None else str(v)) for v in row] for row in data]
                return data, headers
            else:
                data = df.head(limit).fillna("").values.tolist()
                data = [[("" if v is None else str(v)) for v in row] for row in data]
                headers = [f"col{i}" for i in range(len(data[0]) if data else 0)]
                return data, headers

        if ext == ".docx":
            Document = _try_import_docx()
            if Document is None:
                raise RuntimeError("缺少 python-docx。请安装：pip install python-docx")
            doc = Document(path)
            if not doc.tables:
                raise RuntimeError("docx 中未找到表格。")
            table = doc.tables[0]
            rows = []
            for r in table.rows:
                rows.append([cell.text.strip() for cell in r.cells])

            if not rows:
                raise RuntimeError("docx 表格为空。")

            if has_header and len(rows) >= 2:
                headers = rows[0]
                data = rows[1:1+limit]
            else:
                headers = [f"col{i}" for i in range(len(rows[0]))]
                data = rows[:limit]

            return data, headers

        raise RuntimeError("不支持的文件类型。")

    def build_pairs(self):
        path = self.file_path.get().strip()
        if not path:
            messagebox.showwarning("提示", "请先选择文件。")
            return

        try:
            pairs, headers, preview_rows = self._load_pairs_from_file(path)
            if not pairs:
                raise RuntimeError("未读到任何有效的 A-B 行。")
            self.deck.set_pairs(pairs, headers=headers, preview_rows=preview_rows, source=path)
            self.load_info.set(f"✅ 已建立 A-B 对应关系：{self.deck.size()} 对。")
            messagebox.showinfo("成功", f"已建立对应关系：{self.deck.size()} 对。\n\n下一步：去“设置练习”初始化题目顺序。")
        except Exception as e:
            messagebox.showerror("失败", f"建立对应失败：{e}")

    def _load_pairs_from_file(self, path):
        # 读全量（非预览）
        ext = os.path.splitext(path)[1].lower()
        has_header = bool(self.has_header.get())
        a_idx = int(self.col_a.get())
        b_idx = int(self.col_b.get())

        headers = []
        preview_rows = []

        if ext in (".xlsx", ".xls", ".csv", ".tsv", ".txt"):
            pd = _try_import_pandas()
            if pd is None:
                raise RuntimeError("缺少 pandas。请安装：pip install pandas openpyxl")

            if ext in (".xlsx", ".xls"):
                sheet = self.sheet_name.get().strip() or 0
                df = pd.read_excel(path, sheet_name=sheet, header=0 if has_header else None, dtype=str)
            else:
                sep = "\t" if ext == ".tsv" else ("," if ext == ".csv" else None)
                encodings = ["utf-8-sig", "utf-8", "gbk", "gb18030"]
                last_err = None
                df = None
                for enc in encodings:
                    try:
                        df = pd.read_csv(path, sep=sep, header=0 if has_header else None, dtype=str, encoding=enc, engine="python")
                        last_err = None
                        break
                    except Exception as e:
                        last_err = e
                if df is None:
                    raise RuntimeError(f"读取失败：{last_err}")

            if has_header:
                headers = [str(c) for c in df.columns.tolist()]
                values = df.fillna("").values.tolist()
            else:
                values = df.fillna("").values.tolist()
                headers = [f"col{i}" for i in range(len(values[0]) if values else 0)]

            pairs = []
            for row in values:
                a = row[a_idx] if a_idx < len(row) else ""
                b = row[b_idx] if b_idx < len(row) else ""
                pairs.append((a, b))

            preview_rows = values[:30]
            preview_rows = [[("" if v is None else str(v)) for v in row] for row in preview_rows]
            return pairs, headers, preview_rows

        if ext == ".docx":
            Document = _try_import_docx()
            if Document is None:
                raise RuntimeError("缺少 python-docx。请安装：pip install python-docx")
            doc = Document(path)
            if not doc.tables:
                raise RuntimeError("docx 中未找到表格。")
            table = doc.tables[0]
            rows = [[cell.text.strip() for cell in r.cells] for r in table.rows]
            if not rows:
                raise RuntimeError("docx 表格为空。")

            if has_header and len(rows) >= 2:
                headers = rows[0]
                data_rows = rows[1:]
            else:
                headers = [f"col{i}" for i in range(len(rows[0]))]
                data_rows = rows

            pairs = []
            for row in data_rows:
                a = row[a_idx] if a_idx < len(row) else ""
                b = row[b_idx] if b_idx < len(row) else ""
                pairs.append((a, b))

            preview_rows = data_rows[:30]
            return pairs, headers, preview_rows

        raise RuntimeError("不支持的文件类型。")

    # ---------------- Quiz ----------------
    def update_answer_widgets_visibility(self):
        mode = self.quiz_mode.get()
        if mode == "fill":
            self.fill_frame.grid()
            self.mc_frame.grid_remove()
        else:
            self.fill_frame.grid_remove()
            self.mc_frame.grid()
        self._layout_mc_buttons()

    def _layout_mc_buttons(self):
        # 根据 num_choices 显示按钮
        n = int(self.num_choices.get())
        n = max(2, min(8, n))
        # 清空布局
        for btn in self.mc_buttons:
            btn.grid_forget()

        # 2 列布局
        cols = 2
        for i in range(n):
            r = i // cols
            c = i % cols
            self.mc_buttons[i].grid(row=r, column=c, sticky="ew", padx=6, pady=6)
            self.mc_frame.columnconfigure(c, weight=1)

    def reset_quiz(self):
        if self.deck.size() == 0:
            messagebox.showwarning("提示", "还没有建立 A-B 对应关系。请先在“导入表格”里建立。")
            return

        self.update_answer_widgets_visibility()

        self.correct = 0
        self.total = 0
        self.current_index = -1

        self.order = list(range(self.deck.size()))
        if self.shuffle.get():
            random.shuffle(self.order)

        self.score_var.set("得分：0 / 0")
        self.progress_var.set(f"进度：0 / {len(self.order)}")
        self.feedback_var.set("已初始化 ✅ 点击“下一题 ▶️”开始。")
        self.question_text.set("（准备好了）")
        self.fill_entry.delete(0, tk.END)

    def next_question(self):
        if self.deck.size() == 0:
            messagebox.showwarning("提示", "还没有建立 A-B 对应关系。")
            return
        if not self.order:
            self.reset_quiz()
            return

        self.update_answer_widgets_visibility()

        self.current_index += 1
        if self.current_index >= len(self.order):
            self.feedback_var.set(f"🎉 完成！最终得分：{self.correct} / {self.total}")
            self.question_text.set("（已结束）")
            return

        idx = self.order[self.current_index]
        a, b = self.deck.pairs[idx]

        if self.direction.get() == "A2B":
            self.current_x, self.current_y = a, b
        else:
            self.current_x, self.current_y = b, a

        self.question_text.set(self.current_x if self.current_x else "（空）")
        self.feedback_var.set("")

        # prepare options
        if self.quiz_mode.get() == "mc":
            self._prepare_mc_options()

        # update progress
        self.progress_var.set(f"进度：{self.current_index+1} / {len(self.order)}")

        # focus entry for fill
        if self.quiz_mode.get() == "fill":
            self.fill_entry.delete(0, tk.END)
            self.fill_entry.focus_set()

    def _prepare_mc_options(self):
        n = int(self.num_choices.get())
        n = max(2, min(8, n))
        correct = self.current_y

        # decoys from all possible answers
        all_answers = [b if self.direction.get() == "A2B" else a for (a, b) in self.deck.pairs]
        decoys = [x for x in all_answers if x != correct and x != ""]
        random.shuffle(decoys)

        options = [correct]
        for d in decoys:
            if len(options) >= n:
                break
            if d not in options:
                options.append(d)

        # 不足时补空（极端情况）
        while len(options) < n:
            options.append("（无）")

        random.shuffle(options)
        self.current_options = options

        self._layout_mc_buttons()
        for i in range(n):
            self.mc_buttons[i].configure(text=options[i])

    def submit_choice(self, option_index):
        n = int(self.num_choices.get())
        n = max(2, min(8, n))
        if option_index >= n or option_index >= len(self.current_options):
            return

        picked = self.current_options[option_index]
        self.total += 1

        if picked == self.current_y:
            self.correct += 1
            self.feedback_var.set("✅ 正确！")
        else:
            self.feedback_var.set(f"❌ 错了。正确答案：{self.current_y}")

        self.score_var.set(f"得分：{self.correct} / {self.total}")

    def submit_fill(self):
        user_ans = self.fill_entry.get()
        self.total += 1

        a = normalize_text(user_ans, self.case_insensitive.get())
        b = normalize_text(self.current_y, self.case_insensitive.get())

        if a == b and b != "":
            self.correct += 1
            self.feedback_var.set("✅ 正确！")
        else:
            self.feedback_var.set(f"❌ 不对。正确答案：{self.current_y}")

        self.score_var.set(f"得分：{self.correct} / {self.total}")

def main():
    app = ABApp()
    # 让设置变化立即影响UI（特别是模式切换）
    app.quiz_mode.trace_add("write", lambda *_: app.update_answer_widgets_visibility())
    app.num_choices.trace_add("write", lambda *_: app._layout_mc_buttons())
    app.mainloop()

if __name__ == "__main__":
    main()